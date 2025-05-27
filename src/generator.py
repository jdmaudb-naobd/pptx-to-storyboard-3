"""
Generate Word document from processed content
"""

from docx import Document
from docx.shared import Inches
from pathlib import Path
from typing import Dict, List
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_grey(cell, hex_color="D9D9D9"):
    """Set a very light grey background for a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

class StoryboardGenerator:
    def __init__(self, template_path=None):
        if template_path and Path(template_path).exists():
            self.doc = Document(template_path)
        else:
            self.doc = Document()
    
    def create_title_page(self, title: str):
        """Create title page"""
        self.doc.add_heading(title, 0)
        self.doc.add_page_break()
    
    def create_contents_table(self, structure: Dict):
        """Create table of contents"""
        self.doc.add_heading('Table of Contents', 1)
        
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Chapter'
        hdr_cells[1].text = 'Page/Slide'
        
        # Add chapters
        for chapter in structure.get('chapters', []):
            row_cells = table.add_row().cells
            row_cells[0].text = chapter['title']
            row_cells[1].text = str(chapter['slide_number'])
            
            # Add subchapters
            for subchapter in chapter.get('subchapters', []):
                row_cells = table.add_row().cells
                row_cells[0].text = f"  - {subchapter['title']}"
                row_cells[1].text = str(subchapter['slide_number'])
        
        self.doc.add_page_break()
    
    def create_abbreviations_table(self, abbreviations: Dict):
        """Create abbreviations table"""
        if not abbreviations:
            return
            
        self.doc.add_heading('Abbreviations', 1)
        
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Abbreviation'
        hdr_cells[1].text = 'Definition'
        
        # Add abbreviations
        for abbr, definition in sorted(abbreviations.items()):
            row_cells = table.add_row().cells
            row_cells[0].text = abbr
            row_cells[1].text = definition
        
        self.doc.add_page_break()
    
    def create_objectives_section(self, objectives: List[str]):
        """Create learning objectives section"""
        if not objectives:
            return
            
        self.doc.add_heading('Learning Objectives', 1)
        
        for obj in objectives:
            self.doc.add_paragraph(f'• {obj}')
        
        self.doc.add_page_break()
    
    def create_content_table(self, slide_data: Dict, chapter: str = "", subchapter: str = "", references: List[str] = None):
        """Create content table for a slide"""
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        # Table content
        rows_data = [
            ("Chapter", chapter),
            ("Subchapter", subchapter),
            ("Text", "\n".join([t["text"] for t in slide_data.get("texts", [])])),
            ("Media/Images", f"{len(slide_data.get('shapes', []))} images" if slide_data.get('shapes') else "None"),
            ("Visual Details", "See slide"),
            ("Interactivity Details", "None"),
            ("References", "; ".join(references) if references else ""),
            ("Extra Details/Settings", "")
        ]
        
        for idx, (label, content) in enumerate(rows_data):
            row = table.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = content
            set_cell_grey(row.cells[0])  # Set grey background for the label cell

        self.doc.add_paragraph()  # Add spacing
    
    def save(self, output_path: str):
        """Save the document"""
        self.doc.save(output_path)
        print(f"Storyboard saved to {output_path}")

def generate_storyboard(structure: Dict, content: Dict, references: Dict, output_path: str, title: str = "Storyboard"):
    """Generate the storyboard document"""
    generator = StoryboardGenerator()
    generator.create_title_page(title)
    
    # Create Table of Contents
    generator.create_contents_table(structure)
    
    # Create Abbreviations table
    abbreviations = {  # Example abbreviations, replace with actual data
        "AI": "Artificial Intelligence",
        "ML": "Machine Learning",
        "NLP": "Natural Language Processing"
    }
    generator.create_abbreviations_table(abbreviations)
    
    # Create Learning Objectives section
    objectives = [
        "Understand the basics of AI",
        "Learn about machine learning algorithms",
        "Get introduced to natural language processing"
    ]
    generator.create_objectives_section(objectives)
    
    # Add content for each chapter and subchapter
    for chapter in structure.get('chapters', []):
        current_chapter = chapter['title']
        for slide_num in chapter.get('slides', []):
            slide_data = content['slides'][slide_num - 1]
            slide_refs = references.get(slide_num, [])
            generator.create_content_table(slide_data, current_chapter, "", slide_refs)
        for subchapter in chapter.get('subchapters', []):
            current_subchapter = subchapter['title']
            for slide_num in subchapter.get('slides', []):
                slide_data = content['slides'][slide_num - 1]
                slide_refs = references.get(slide_num, [])
                generator.create_content_table(slide_data, current_chapter, current_subchapter, slide_refs)
    
    # Save the document
    generator.save(output_path)
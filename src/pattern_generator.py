"""
Pattern-based storyboard generator that uses learned patterns
"""

import json
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from collections import defaultdict

from src.extractor import SimpleExtractor
from src.medical_processor import MedicalContentProcessor
from src.generator import StoryboardGenerator
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from src.utils import sanitize_text


class PatternBasedGenerator:
    """Generate storyboards using learned patterns"""
    
    def __init__(self, patterns_file: str = "learned_patterns.json"):
        """Initialize with learned patterns"""
        self.patterns = self._load_patterns(patterns_file)
        self.medical_processor = MedicalContentProcessor()
        
    def _load_patterns(self, patterns_file: str) -> Dict:
        """Load learned patterns"""
        try:
            with open(patterns_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            print(f"Warning: {patterns_file} not found. Using defaults.")
            return self._get_default_patterns()
    
    def _get_default_patterns(self) -> Dict:
        """Return default patterns if file not found"""
        return {
            'transformation_rules': {
                'omit_slide_types': ['disclosure', 'title', 'references'],
                'combine_slide_types': ['content', 'patient_case'],
                'avg_slides_per_segment': 2.5
            },
            'structure_template': {
                'standard_sequence': [
                    'Welcome',
                    'Learning objectives',
                    'Introduction',
                    'Main Content',
                    'Summary',
                    'Thank you'
                ]
            }
        }
    
    def generate_storyboard(self, pptx_path: str, output_path: str = None) -> str:
        """Generate a storyboard using learned patterns"""
        pptx_path = Path(pptx_path)
        if not output_path:
            output_path = pptx_path.with_suffix('_storyboard.docx')
        
        print(f"\n🎯 Pattern-Based Storyboard Generation")
        print("=" * 60)
        print(f"Input: {pptx_path.name}")
        
        # Step 1: Extract and analyze PowerPoint
        print("\n1. Extracting content from PowerPoint...")
        extractor = SimpleExtractor(pptx_path)
        pptx_content = extractor.extract_all_content()
        pptx_structure = self.medical_processor.identify_structure(pptx_content)
        
        # Step 2: Apply transformation patterns
        print("\n2. Applying learned transformation patterns...")
        transformed_structure = self._apply_transformation_patterns(
            pptx_content, pptx_structure
        )
        
        # Step 3: Extract key information
        print("\n3. Extracting abbreviations and objectives...")
        abbreviations = self.medical_processor.extract_abbreviations(pptx_content)
        objectives = self._extract_or_generate_objectives(pptx_content, pptx_structure)
        references = self.medical_processor.extract_references(pptx_content)
        
        # Step 4: Create storyboard structure
        print("\n4. Creating storyboard structure...")
        storyboard_structure = self._create_storyboard_structure(
            transformed_structure, pptx_content
        )
        
        # Step 5: Generate document
        print("\n5. Generating storyboard document...")
        self._generate_document(
            output_path,
            storyboard_structure,
            abbreviations,
            objectives,
            references,
            pptx_path.stem
        )
        
        print(f"\n✅ Storyboard generated: {output_path}")
        print("=" * 60)
        
        return str(output_path)
    
    def _apply_transformation_patterns(self, content: Dict, structure: Dict) -> Dict:
        """Apply learned transformation rules"""
        transformed = {
            'chapters': [],
            'omitted_slides': [],
            'combined_slides': []
        }
        
        rules = self.patterns['transformation_rules']
        omit_types = rules['omit_slide_types']
        combine_types = rules['combine_slide_types']
        
        # Track which slides to include
        slides_to_process = []
        slides_by_type = defaultdict(list)
        
        for slide in content['slides']:
            slide_num = slide['slide_number']
            slide_type = structure['slide_types'].get(slide_num, 'content')
            
            # Check if should omit
            if slide_type in omit_types:
                transformed['omitted_slides'].append({
                    'slide': slide_num,
                    'type': slide_type,
                    'reason': 'Pattern: commonly omitted'
                })
            else:
                slides_to_process.append(slide)
                slides_by_type[slide_type].append(slide)
        
        print(f"   - Omitted {len(transformed['omitted_slides'])} slides")
        print(f"   - Processing {len(slides_to_process)} slides")
        
        # Apply slide type priorities
        priority_rules = rules.get('slide_type_priority', {})
        
        # Group slides for combination
        slide_groups = self._group_slides_for_combination(
            slides_to_process, structure, combine_types
        )
        
        print(f"   - Created {len(slide_groups)} content groups")
        
        transformed['slide_groups'] = slide_groups
        return transformed
    
    def _group_slides_for_combination(self, slides: List[Dict], 
                                    structure: Dict, 
                                    combine_types: List[str]) -> List[Dict]:
        """Group slides based on combination patterns"""
        groups = []
        current_group = None
        
        for slide in slides:
            slide_num = slide['slide_number']
            slide_type = structure['slide_types'].get(slide_num, 'content')
            
            # Check if this starts a new group
            if self._should_start_new_group(slide, slide_type, current_group):
                if current_group:
                    groups.append(current_group)
                current_group = {
                    'slides': [slide],
                    'primary_type': slide_type,
                    'slide_numbers': [slide_num]
                }
            else:
                # Add to current group
                if current_group:
                    current_group['slides'].append(slide)
                    current_group['slide_numbers'].append(slide_num)
                else:
                    current_group = {
                        'slides': [slide],
                        'primary_type': slide_type,
                        'slide_numbers': [slide_num]
                    }
        
        # Don't forget the last group
        if current_group:
            groups.append(current_group)
        
        return groups
    
    def _should_start_new_group(self, slide: Dict, slide_type: str, 
                               current_group: Optional[Dict]) -> bool:
        """Determine if slide should start a new group"""
        if not current_group:
            return True
        
        # Always start new group for certain types
        always_separate = ['objectives', 'conclusion', 'questions']
        if slide_type in always_separate:
            return True
        
        # Check if we've reached combination limit
        avg_slides = self.patterns['transformation_rules']['avg_slides_per_segment']
        if len(current_group['slides']) >= avg_slides:
            return True
        
        # Check word count threshold
        current_words = sum(
            len(' '.join(t['text'] for t in s['texts']).split())
            for s in current_group['slides']
        )
        
        if current_words > self.patterns['content_rules']['split_threshold']:
            return True
        
        return False
    
    def _extract_or_generate_objectives(self, content: Dict, 
                                      structure: Dict) -> List[str]:
        """Extract objectives or generate from content"""
        # First try to extract
        objectives = self.medical_processor.extract_objectives(content)
        
        if not objectives:
            # Generate based on content
            print("   - No objectives found, generating from content...")
            objectives = self._generate_objectives_from_content(content, structure)
        
        return objectives
    
    def _generate_objectives_from_content(self, content: Dict, 
                                        structure: Dict) -> List[str]:
        """Generate learning objectives from content"""
        objectives = []
        
        # Look for key topics in content
        key_topics = set()
        for slide in content['slides']:
            slide_type = structure['slide_types'].get(slide['slide_number'], '')
            if slide_type in ['clinical_data', 'treatment', 'patient_case']:
                # Extract key concepts
                text = ' '.join(t['text'] for t in slide['texts'])
                # Simple extraction of important terms
                if 'treatment' in text.lower():
                    key_topics.add('treatment options')
                if 'clinical' in text.lower():
                    key_topics.add('clinical evidence')
                if 'patient' in text.lower():
                    key_topics.add('patient management')
        
        # Generate objectives
        if 'treatment options' in key_topics:
            objectives.append("Understand the treatment options and their clinical applications")
        if 'clinical evidence' in key_topics:
            objectives.append("Review the clinical evidence and study outcomes")
        if 'patient management' in key_topics:
            objectives.append("Apply patient management strategies in clinical practice")
        
        # Add default if none found
        if not objectives:
            objectives = [
                "Understand the key concepts presented in this module",
                "Apply the learning to clinical practice",
                "Evaluate treatment options based on evidence"
            ]
        
        return objectives
    
    def _create_storyboard_structure(self, transformed: Dict, 
                                   content: Dict) -> Dict:
        """Create the storyboard structure using patterns"""
        structure = {
            'chapters': []
        }
        
        # Use the standard sequence from patterns
        standard_sequence = self.patterns['structure_template']['standard_sequence']
        
        # Map slide groups to chapters
        slide_groups = transformed.get('slide_groups', [])
        group_index = 0
        
        for chapter_name in standard_sequence:
            chapter = {
                'title': chapter_name,
                'segments': []
            }
            
            # Determine how many segments for this chapter
            if chapter_name == 'Main Content':
                # Main content gets most of the slide groups
                remaining_groups = len(slide_groups) - group_index - 2  # Save some for summary
                segments_to_add = max(1, remaining_groups)
                
                for i in range(segments_to_add):
                    if group_index < len(slide_groups):
                        segment = self._create_segment_from_group(
                            slide_groups[group_index], 
                            chapter_name,
                            f"Section {i+1}"
                        )
                        chapter['segments'].append(segment)
                        group_index += 1
            
            elif chapter_name in ['Welcome', 'Meet the experts']:
                # These might not have slides, create placeholder
                chapter['segments'].append({
                    'chapter': chapter_name,
                    'subchapter': '',
                    'content': self._get_placeholder_content(chapter_name),
                    'is_placeholder': True
                })
            
            elif chapter_name == 'Learning objectives':
                # Special handling for objectives
                chapter['segments'].append({
                    'chapter': chapter_name,
                    'subchapter': '',
                    'content': 'See objectives section above',
                    'is_objectives': True
                })
            
            elif chapter_name in ['Pre-assessment questions', 'Post-assessment questions']:
                # Create question placeholders
                for i in range(3):  # 3 questions each
                    chapter['segments'].append({
                        'chapter': chapter_name,
                        'subchapter': f'Question {i+1}',
                        'content': self._create_question_placeholder(i+1),
                        'is_question': True
                    })
            
            else:
                # For other chapters, use next available group if exists
                if group_index < len(slide_groups):
                    segment = self._create_segment_from_group(
                        slide_groups[group_index], 
                        chapter_name,
                        ''
                    )
                    chapter['segments'].append(segment)
                    group_index += 1
                else:
                    # Create placeholder
                    chapter['segments'].append({
                        'chapter': chapter_name,
                        'subchapter': '',
                        'content': self._get_placeholder_content(chapter_name),
                        'is_placeholder': True
                    })
            
            structure['chapters'].append(chapter)
        
        return structure
    
    def _create_segment_from_group(self, group: Dict, 
                                  chapter: str, 
                                  subchapter: str) -> Dict:
        """Create a segment from a slide group"""
        # Combine text from all slides in group
        all_texts = []
        for slide in group['slides']:
            for text_item in slide['texts']:
                text = text_item['text'].strip()
                if text:
                    all_texts.append(text)
        
        # Apply content rules
        content = self._format_content_by_rules(all_texts)
        
        segment = {
            'chapter': chapter,
            'subchapter': subchapter,
            'content': content,
            'source_slides': group['slide_numbers'],
            'images': sum(len(s.get('shapes', [])) for s in group['slides'])
        }
        
        return segment
    
    def _format_content_by_rules(self, texts: List[str]) -> str:
        """Format content based on learned rules"""
        if not texts:
            return ""
        
        rules = self.patterns['content_rules']
        
        # Count potential bullet points
        bullet_count = sum(1 for t in texts if len(t) < 100 and not t.endswith('.'))
        
        if bullet_count >= rules['bullet_threshold']:
            # Format as bullets
            formatted = []
            for text in texts:
                if text and not text[0].isupper():
                    text = text[0].upper() + text[1:]
                formatted.append(f"• {text}")
            return '\n'.join(formatted)
        else:
            # Format as paragraphs
            return '\n\n'.join(texts)
    
    def _get_placeholder_content(self, chapter_name: str) -> str:
        """Get placeholder content for chapters without slides"""
        placeholders = {
            'Welcome': '[Welcome message to be added]',
            'Meet the experts': '[Expert bio and credentials to be added]',
            'Introduction': '[Introduction content to be developed]',
            'Summary': '[Summary of key points to be added]',
            'Thank you': '[Thank you message and contact information]'
        }
        return placeholders.get(chapter_name, f'[{chapter_name} content to be added]')
    
    def _create_question_placeholder(self, question_num: int) -> Dict:
        """Create a question placeholder"""
        return {
            'question': f'[Question {question_num} to be developed]',
            'answers': [
                '[Answer option A]',
                '[Answer option B]',
                '[Answer option C]',
                '[Answer option D]'
            ],
            'correct': 0,
            'feedback': '[Feedback to be added]'
        }
    
    def _generate_document(self, output_path: str, structure: Dict,
                          abbreviations: Dict, objectives: List[str],
                          references: Dict, title: str):
        """Generate the Word document"""
        generator = StoryboardGenerator()
        
        # Title page
        generator.create_title_page(f"{title} - eLearning Storyboard")
        
        # Table of contents
        toc_structure = {
            'chapters': [
                {
                    'title': ch['title'],
                    'slide_number': i+1,
                    'subchapters': [
                        {'title': seg['subchapter'], 'slide_number': i+1}
                        for seg in ch['segments'] if seg.get('subchapter')
                    ]
                }
                for i, ch in enumerate(structure['chapters'])
            ]
        }
        generator.create_contents_table(toc_structure)
        
        # Abbreviations
        if abbreviations:
            generator.create_abbreviations_table(abbreviations)
        
        # Objectives
        if objectives:
            generator.create_objectives_section(objectives)
        
        # Content segments
        for chapter in structure['chapters']:
            generator.doc.add_heading(chapter['title'], 1)
            
            for segment in chapter['segments']:
                if segment.get('is_question'):
                    # Create question table
                    self._create_question_table(generator, segment)
                elif segment.get('is_objectives'):
                    # Skip - already in objectives section
                    continue
                else:
                    # Create content table
                    self._create_content_segment_table(
                        generator, segment, abbreviations, references
                    )
        
        # Save
        generator.save(output_path)
    
    def _create_content_segment_table(self, generator: StoryboardGenerator,
                                    segment: Dict, abbreviations: Dict,
                                    references: Dict):
        """Create a content segment table"""
        # Prepare slide references
        slide_refs = []
        if 'source_slides' in segment:
            for slide_num in segment['source_slides']:
                if slide_num in references:
                    slide_refs.extend(references[slide_num])
        
        # Create the table
        slide_data = {
            'texts': [{'text': segment['content'], 'is_title': False}],
            'shapes': [{'type': 'image'}] * segment.get('images', 0)
        }
        
        generator.create_content_table(
            slide_data,
            segment['chapter'],
            segment.get('subchapter', ''),
            slide_refs,
            abbreviations
        )
    
    def _create_question_table(self, generator: StoryboardGenerator,
                             question_data: Dict):
        """Create a question table"""
        table = generator.doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Add question content
        rows_data = [
            ("Chapter", question_data.get('chapter', '')),
            ("Subchapter", question_data.get('subchapter', '')),
            ("Text", question_data.get('content', {}).get('question', '[Question text]')),
            ("Answers (correct in green)", '\n'.join(
                question_data.get('content', {}).get('answers', ['[Answers]'])
            )),
            ("Answer related feedback", '[Feedback for each answer]'),
            ("Solution", '[Explanation of correct answer]')
        ]
        
        for idx, (label, content) in enumerate(rows_data):
            row = table.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = content
            
            # Style the label cell (using the method from generator.py)
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            cell = row.cells[0]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), "D9D9D9")
            tcPr.append(shd)
        
        generator.doc.add_paragraph()  # Add spacing
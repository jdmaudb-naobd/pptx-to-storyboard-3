"""
Simplified analyzer that properly reads Word documents
"""

import os
import json
from pathlib import Path
from typing import Dict, List, Tuple
from collections import defaultdict, Counter
from docx import Document

from src.extractor import SimpleExtractor
from src.medical_processor import MedicalContentProcessor


class SimpleDocumentAnalyzer:
    """Simple analyzer that reads documents properly"""
    
    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        
    def extract_structure(self) -> Dict:
        """Extract structure using python-docx properly"""
        structure = {
            'chapters': [],
            'abbreviations': {},
            'segments': []
        }
        
        current_chapter = None
        current_subchapter = None
        
        # Read through all elements
        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                # Get paragraph properly
                for paragraph in self.doc.paragraphs:
                    if paragraph._element == element:
                        text = paragraph.text.strip()
                        style = paragraph.style.name if paragraph.style else 'Normal'
                        
                        if text:  # Only process non-empty
                            if style == 'Heading 1':
                                current_chapter = text
                                structure['chapters'].append({
                                    'title': text,
                                    'subchapters': []
                                })
                            elif style == 'AX Subhead':
                                current_subchapter = text
                                if structure['chapters']:
                                    structure['chapters'][-1]['subchapters'].append(text)
                        break
                        
            elif element.tag.endswith('tbl'):
                # Get table properly
                for table in self.doc.tables:
                    if table._element == element:
                        # Analyze table
                        if self._is_segment_table(table):
                            segment = self._extract_segment_info(table)
                            if current_chapter:
                                segment['chapter'] = current_chapter
                            if current_subchapter:
                                segment['subchapter'] = current_subchapter
                            structure['segments'].append(segment)
                        elif self._is_abbreviation_table(table):
                            abbrevs = self._extract_abbreviations(table)
                            structure['abbreviations'].update(abbrevs)
                        break
        
        return structure
    
    def _is_segment_table(self, table) -> bool:
        """Check if table is a content segment"""
        if len(table.rows) >= 5 and len(table.columns) in [2, 3, 4]:
            # Check first column for segment keywords
            first_col_text = ' '.join(cell.text.lower() for row in table.rows for cell in [row.cells[0]])
            keywords = ['chapter', 'subchapter', 'text', 'visual', 'interactivity', 'reference']
            return sum(1 for kw in keywords if kw in first_col_text) >= 3
        return False
    
    def _is_abbreviation_table(self, table) -> bool:
        """Check if table is abbreviations"""
        if len(table.columns) == 2:
            # Check header
            if table.rows:
                header_text = table.rows[0].cells[0].text.lower() + table.rows[0].cells[1].text.lower()
                return 'abbreviation' in header_text or 'definition' in header_text
        return False
    
    def _extract_segment_info(self, table) -> Dict:
        """Extract info from segment table"""
        segment = {}
        for row in table.rows:
            if len(row.cells) >= 2:
                label = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                
                if 'chapter' in label:
                    segment['chapter'] = value
                elif 'subchapter' in label:
                    segment['subchapter'] = value  
                elif 'text' in label:
                    segment['text'] = value
                elif 'reference' in label:
                    segment['references'] = value
        return segment
    
    def _extract_abbreviations(self, table) -> Dict:
        """Extract abbreviations from table"""
        abbrevs = {}
        for i, row in enumerate(table.rows):
            if i == 0:  # Skip header
                continue
            if len(row.cells) >= 2:
                abbr = row.cells[0].text.strip()
                defn = row.cells[1].text.strip()
                if abbr and defn:
                    abbrevs[abbr] = defn
        return abbrevs


class SimplifiedAnalyzer:
    """Simplified pattern analyzer"""
    
    def __init__(self, examples_dir: str = "examples"):
        self.examples_dir = Path(examples_dir)
        
    def analyze_all_examples(self) -> Dict:
        """Analyze with simplified approach"""
        print("🔍 Running Simplified Analysis...")
        print("=" * 60)
        
        input_dir = self.examples_dir / "input"
        output_dir = self.examples_dir / "output"
        
        projects = [d.name for d in input_dir.iterdir() if d.is_dir()]
        print(f"Found {len(projects)} projects")
        
        # Collect patterns
        all_patterns = {
            'chapter_sequences': [],
            'slide_counts': [],
            'common_chapters': Counter(),
            'slide_to_chapter_ratio': []
        }
        
        for project in projects[:5]:  # Analyze first 5 for speed
            print(f"\nAnalyzing {project}...")
            
            # Get files
            input_files = list((input_dir / project).glob("*.pptx"))
            output_files = list((output_dir / project).glob("*.docx"))
            
            for input_file in input_files:
                # Find matching output
                matching_output = None
                for output_file in output_files:
                    if input_file.stem == output_file.stem:
                        matching_output = output_file
                        break
                
                if matching_output:
                    # Extract from PowerPoint
                    extractor = SimpleExtractor(input_file)
                    pptx_content = extractor.extract_all_content()
                    slide_count = len(pptx_content['slides'])
                    
                    # Extract from Word
                    doc_analyzer = SimpleDocumentAnalyzer(str(matching_output))
                    doc_structure = doc_analyzer.extract_structure()
                    
                    # Collect patterns
                    all_patterns['slide_counts'].append(slide_count)
                    
                    # Chapter sequence
                    chapter_names = [ch['title'] for ch in doc_structure['chapters']]
                    all_patterns['chapter_sequences'].append(chapter_names)
                    
                    # Count chapter occurrences
                    for chapter in chapter_names:
                        all_patterns['common_chapters'][chapter] += 1
                    
                    # Ratio
                    if doc_structure['segments']:
                        all_patterns['slide_to_chapter_ratio'].append(
                            slide_count / len(doc_structure['segments'])
                        )
        
        # Generate simple patterns
        patterns = self._generate_patterns(all_patterns)
        
        # Save patterns
        with open("learned_patterns.json", "w", encoding="utf-8") as f:
            json.dump(patterns, f, indent=2)
        
        print("\n✅ Analysis complete!")
        print(f"📄 Patterns saved to: learned_patterns.json")
        
        return patterns
    
    def _generate_patterns(self, all_patterns: Dict) -> Dict:
        """Generate usable patterns from analysis"""
        # Most common chapters
        common_chapters = [ch for ch, count in all_patterns['common_chapters'].most_common(15)]
        
        # Average slide count
        avg_slides = sum(all_patterns['slide_counts']) / len(all_patterns['slide_counts']) if all_patterns['slide_counts'] else 30
        
        # Common chapter sequence (simplified)
        standard_sequence = []
        for chapters in all_patterns['chapter_sequences']:
            if len(chapters) > len(standard_sequence):
                standard_sequence = chapters
        
        patterns = {
            'transformation_rules': {
                'omit_slide_types': ['disclosure', 'title', 'references'],
                'combine_slide_types': ['content', 'patient_case'],
                'avg_slides_per_segment': 2.5
            },
            'structure_template': {
                'common_chapters': common_chapters,
                'standard_sequence': standard_sequence if standard_sequence else [
                    'Welcome',
                    'Meet the experts', 
                    'Learning objectives',
                    'Pre-assessment questions',
                    'Introduction',
                    'Main Content',
                    'Summary',
                    'Post-assessment questions',
                    'Thank you'
                ]
            },
            'statistics': {
                'avg_slide_count': avg_slides,
                'avg_chapters': len(common_chapters)
            }
        }
        
        return patterns
